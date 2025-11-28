
from pathlib import Path
from uuid import uuid4

import pytest
from fastapi.testclient import TestClient
from docx import Document as DocxDocument

from main import app, STORAGE_DIR


@pytest.fixture
def temp_storage(tmp_path, monkeypatch):
    """Фикстура для изоляции тестов - использует временную директорию."""
    monkeypatch.setattr("main.STORAGE_DIR", Path(tmp_path))
    from main import app as fresh_app
    client = TestClient(fresh_app)
    return client, tmp_path


class TestDocumentServiceComponent:
    """КОМПОНЕНТНЫЕ ТЕСТЫ: Document Service (тестируется компонент как единое целое)"""

    def test_component_generate_pdf_document(self, temp_storage):
        """Компонентный тест: генерация PDF документа."""
        client, storage_path = temp_storage

        payload = {
            "template_code": "ANALYSIS_RESULT",
            "format": "PDF",
            "data": {
                "analysis": {
                    "patient_name": "Иванов Иван",
                    "hemoglobin": 145,
                    "conclusion": "Показатели в норме",
                    "date": "2025-01-15",
                }
            },
        }

        response = client.post("/api/v1/documents", json=payload)

        assert response.status_code == 200
        body = response.json()
        assert body["format"] == "PDF"
        assert body["file_name"].endswith(".pdf")

        # Проверяем создание файла в файловой системе (реальная ФС, не мок)
        file_path = storage_path / body["file_name"]
        assert file_path.exists()
        assert file_path.stat().st_size > 0

        # Проверяем валидность PDF
        with open(file_path, "rb") as f:
            assert f.read().startswith(b"%PDF")

    def test_component_generate_docx_document(self, temp_storage):
        """Компонентный тест: генерация DOCX документа."""
        client, storage_path = temp_storage

        payload = {
            "template_code": "ANALYSIS_RESULT",
            "format": "DOCX",
            "data": {
                "analysis": {
                    "patient_name": "Петрова Мария",
                    "hemoglobin": 120,
                    "conclusion": "Требуется наблюдение",
                }
            },
        }

        response = client.post("/api/v1/documents", json=payload)

        assert response.status_code == 200
        body = response.json()
        assert body["format"] == "DOCX"
        assert body["file_name"].endswith(".docx")

        # Проверяем создание файла
        file_path = storage_path / body["file_name"]
        assert file_path.exists()

        # Проверяем содержимое через python-docx
        doc = DocxDocument(str(file_path))
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "РЕЗУЛЬТАТ ЛАБОРАТОРНОГО АНАЛИЗА" in text_content
        assert "Петрова Мария" in text_content
        assert "Гемоглобин: 120" in text_content

    def test_component_download_document(self, temp_storage):
        """Компонентный тест: скачивание документа (полный цикл)."""
        client, storage_path = temp_storage

        # Step 1: Создание документа
        payload = {
            "template_code": "ANALYSIS_RESULT",
            "format": "PDF",
            "data": {
                "analysis": {
                    "patient_name": "Сидоров Сидор",
                    "hemoglobin": 130,
                    "conclusion": "Норма",
                }
            },
        }

        create_resp = client.post("/api/v1/documents", json=payload)
        assert create_resp.status_code == 200
        document_id = create_resp.json()["document_id"]
        file_name = create_resp.json()["file_name"]

        # Step 2: Проверка файла на диске
        file_path = storage_path / file_name
        assert file_path.exists()
        file_size = file_path.stat().st_size

        # Step 3: Скачивание через API
        download_resp = client.get(f"/api/v1/documents/{document_id}/content")
        assert download_resp.status_code == 200
        assert len(download_resp.content) == file_size
        assert download_resp.headers["content-type"] == "application/pdf"
        assert download_resp.content.startswith(b"%PDF")

    def test_component_download_docx_with_content_check(self, temp_storage):
        """Компонентный тест: скачивание DOCX с проверкой содержимого."""
        client, storage_path = temp_storage

        # Создаём DOCX
        payload = {
            "template_code": "ANALYSIS_RESULT",
            "format": "DOCX",
            "data": {
                "analysis": {
                    "patient_name": "Тестовый Пациент",
                    "hemoglobin": 140,
                    "conclusion": "Всё в порядке",
                    "leukocytes": 6.0,
                    "erythrocytes": 5.0,
                }
            },
        }

        create_resp = client.post("/api/v1/documents", json=payload)
        document_id = create_resp.json()["document_id"]

        # Скачиваем и проверяем содержимое
        download_resp = client.get(f"/api/v1/documents/{document_id}/content")
        assert download_resp.status_code == 200

        temp_file = storage_path / "check.docx"
        with open(temp_file, "wb") as f:
            f.write(download_resp.content)

        doc = DocxDocument(str(temp_file))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        assert "РЕЗУЛЬТАТ ЛАБОРАТОРНОГО АНАЛИЗА" in full_text
        assert "Тестовый Пациент" in full_text
        assert "Гемоглобин: 140" in full_text
        assert "Дополнительные параметры:" in full_text

    def test_component_custom_template(self, temp_storage):
        """Компонентный тест: генерация документа с кастомным шаблоном."""
        client, storage_path = temp_storage

        payload = {
            "template_code": "CUSTOM_REPORT",
            "format": "DOCX",
            "data": {
                "custom_field": "custom_value",
                "number": 42,
            },
        }

        response = client.post("/api/v1/documents", json=payload)
        assert response.status_code == 200

        document_id = response.json()["document_id"]
        download_resp = client.get(f"/api/v1/documents/{document_id}/content")
        assert download_resp.status_code == 200

        # Проверяем содержимое
        temp_file = storage_path / "custom.docx"
        with open(temp_file, "wb") as f:
            f.write(download_resp.content)

        doc = DocxDocument(str(temp_file))
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Шаблон: CUSTOM_REPORT" in text
        assert "custom_field" in text or '"custom_field"' in text

    def test_component_analysis_template_excludes_id(self, temp_storage):
        """Компонентный тест: шаблон ANALYSIS_RESULT не включает служебные поля."""
        client, storage_path = temp_storage

        payload = {
            "template_code": "ANALYSIS_RESULT",
            "format": "DOCX",
            "data": {
                "analysis": {
                    "id": "secret-id-12345",
                    "patient_name": "Пациент",
                    "hemoglobin": 130,
                    "conclusion": "Норма",
                }
            },
        }

        create_resp = client.post("/api/v1/documents", json=payload)
        document_id = create_resp.json()["document_id"]

        download_resp = client.get(f"/api/v1/documents/{document_id}/content")
        temp_file = storage_path / "id_check.docx"
        with open(temp_file, "wb") as f:
            f.write(download_resp.content)

        doc = DocxDocument(str(temp_file))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # ID не должен быть в документе
        assert "secret-id-12345" not in full_text

    def test_component_download_nonexistent_document(self, temp_storage):
        """Компонентный тест: обработка ошибки при скачивании несуществующего документа."""
        client, _ = temp_storage

        fake_id = str(uuid4())
        response = client.get(f"/api/v1/documents/{fake_id}/content")

        assert response.status_code == 404
        assert "not found" in response.json()["detail"].lower()

    def test_component_full_workflow(self, temp_storage):
        """Компонентный тест: полный цикл работы компонента."""
        client, storage_path = temp_storage

        # Step 1: Создание PDF
        payload = {
            "template_code": "ANALYSIS_RESULT",
            "format": "PDF",
            "data": {
                "analysis": {
                    "patient_name": "Полный Цикл",
                    "hemoglobin": 135,
                    "conclusion": "Тест пройден",
                }
            },
        }

        create_resp = client.post("/api/v1/documents", json=payload)
        assert create_resp.status_code == 200

        document_id = create_resp.json()["document_id"]
        file_name = create_resp.json()["file_name"]

        # Step 2: Проверка файла на диске
        file_path = storage_path / file_name
        assert file_path.exists()

        # Step 3: Скачивание
        download_resp = client.get(f"/api/v1/documents/{document_id}/content")
        assert download_resp.status_code == 200
        assert download_resp.content.startswith(b"%PDF")

        # Step 4: Создание DOCX с теми же данными
        docx_payload = payload.copy()
        docx_payload["format"] = "DOCX"
        docx_resp = client.post("/api/v1/documents", json=docx_payload)
        assert docx_resp.status_code == 200

        docx_id = docx_resp.json()["document_id"]
        docx_file = storage_path / f"{docx_id}.docx"
        assert docx_file.exists()

        # Step 5: Проверка содержимого DOCX
        doc = DocxDocument(str(docx_file))
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Полный Цикл" in text
        assert "Гемоглобин: 135" in text
